"""
Gerador do contrato DOCX formatado juridicamente.
Programa GEO IPOG — Pós-Graduações em Psicologia (reemissão 12-05-2026).
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

OUT_PATH = Path(r"C:\Sandyboxclaude\geo-ipog\docs\governance\contratos\contrato-geo-ipog-reemissao-12-05-2026.docx")

doc = Document()

# Margens forenses (3cm laterais, 2.5cm vertical) e papel A4
section = doc.sections[0]
section.page_height = Cm(29.7)
section.page_width = Cm(21.0)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(3.0)
section.right_margin = Cm(3.0)

# Estilo base
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)
style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

def add_paragraph(text, *, align=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=False, size=12,
                  space_before=0, space_after=6, line_spacing=1.5,
                  first_line_indent=None, all_caps=False, center=False):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    if first_line_indent is not None:
        pf.first_line_indent = Cm(first_line_indent)
    run = p.add_run(text.upper() if all_caps else text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    return p

def add_clause_title(num, text):
    add_paragraph(f"CLÁUSULA {num} — {text}", align=WD_ALIGN_PARAGRAPH.LEFT,
                  bold=True, size=12, space_before=18, space_after=8, line_spacing=1.5)

def add_clause_body(text, *, indent=0):
    add_paragraph(text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12,
                  space_before=0, space_after=6, line_spacing=1.5,
                  first_line_indent=indent if indent else None)

def add_signature_line(label):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(24)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0
    run = p.add_run("_" * 60)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf2 = p2.paragraph_format
    pf2.space_before = Pt(0)
    pf2.space_after = Pt(4)
    pf2.line_spacing = 1.15
    r2 = p2.add_run(label)
    r2.font.name = "Times New Roman"
    r2.font.size = Pt(11)

# ========== TÍTULO ==========
add_paragraph("CONTRATO DE PRESTAÇÃO DE SERVIÇOS DE CONSULTORIA, EXECUÇÃO E TREINAMENTO EM GENERATIVE ENGINE OPTIMIZATION (GEO)",
              center=True, bold=True, size=13, space_after=4, line_spacing=1.2, all_caps=True)
add_paragraph("Programa GEO IPOG — Pós-Graduações em Psicologia",
              center=True, bold=True, size=12, space_after=4, line_spacing=1.2)
add_paragraph(
    "Escopo amplo cobrindo as cinco modalidades canônicas de pós-graduação em Psicologia: "
    "Especialização Lato Sensu, MBA correlato, Mestrado Profissional, Especialização Clínica "
    "certificada por Conselhos profissionais e Formações Híbridas.",
    center=True, size=11, space_after=4, line_spacing=1.2)
add_paragraph(
    "Reemissão de 12 de maio de 2026 — substitui integralmente a versão anterior que tratava do "
    "recorte vertical de MBA Online de Psicologia como objeto único. A presente reemissão formaliza "
    "o escopo amplo do programa, sem alteração de valor, prazo, partes ou demais condições econômicas.",
    center=True, size=10, space_after=18, line_spacing=1.2)

# ========== QUALIFICAÇÃO DAS PARTES ==========
add_paragraph("QUALIFICAÇÃO DAS PARTES", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT,
              size=12, space_before=6, space_after=10, line_spacing=1.5)

add_paragraph("CONTRATANTE", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT,
              size=12, space_before=0, space_after=4, line_spacing=1.5)
add_clause_body(
    "IPOG — INSTITUTO DE PÓS-GRADUAÇÃO & GRADUAÇÃO LTDA, pessoa jurídica de direito privado, "
    "inscrita no CNPJ/MF sob o nº 04.688.977/0001-02, com sede na Rua T-55, nº 713, Quadra 105, "
    "Lote 01-E, Setor Bueno, Goiânia/GO, CEP 74.215-170, neste ato representada na forma de seu "
    "contrato social por seu Chief Executive Officer (CEO), Ronan Maia, brasileiro, portador da "
    "Cédula de Identidade RG nº _______________________ e inscrito no CPF/MF sob o nº "
    "_______________________, doravante denominada simplesmente CONTRATANTE ou IPOG.")

add_paragraph("CONTRATADA", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT,
              size=12, space_before=8, space_after=4, line_spacing=1.5)
add_clause_body(
    "BRGEO LTDA, nome fantasia BRASIL GEO, pessoa jurídica de direito privado, inscrita no "
    "CNPJ/MF sob o nº 66.051.295/0001-33, com sede na Rua do Parque, nº 361, Quadra 145, Lote "
    "Área 3, Bairro Jardim Atlântico (Hub Cerrado), Goiânia/GO, CEP 74.343-245, neste ato "
    "representada na forma de seu contrato social por seu sócio-administrador, Alexandre "
    "Caramaschi, brasileiro, CEO da Brasil GEO, ex-CMO da Semantix (Nasdaq) e cofundador da AI "
    "Brasil, doravante denominada simplesmente CONTRATADA ou BRASIL GEO.")

add_clause_body(
    "CONTRATANTE e CONTRATADA, em conjunto denominadas Partes e, individualmente, Parte, têm "
    "entre si justo e contratado o presente Contrato de Prestação de Serviços (\"Contrato\"), "
    "que se regerá pelas cláusulas e condições a seguir.")

# ========== CLÁUSULA 1ª — OBJETO ==========
add_clause_title("1ª", "OBJETO")
add_clause_body(
    "1.1. O presente Contrato tem por objeto a prestação, pela CONTRATADA, de serviços continuados "
    "de Generative Engine Optimization (GEO) voltados ao posicionamento da CONTRATANTE em motores "
    "generativos de inteligência artificial (ChatGPT, Claude, Gemini, Perplexity, Copilot e "
    "congêneres), com foco prioritário no portfólio de Pós-Graduações em Psicologia da CONTRATANTE, "
    "abrangendo de forma integrada todas as modalidades regulamentadas de pós-graduação em "
    "Psicologia ofertadas pela CONTRATANTE no presente e ao longo da vigência do Contrato.")
add_clause_body(
    "1.2. As modalidades canônicas que integram o escopo amplo do programa, na forma do reframe "
    "canônico aplicado em 12 de maio de 2026 e formalizado no repositório oficial de governança da "
    "CONTRATADA, são:")
add_clause_body(
    "a) Especialização Lato Sensu em Psicologia (carga horária mínima de 360 horas, autorizada "
    "pelo MEC), cobrindo Neuropsicologia, Avaliação Psicológica, Psicologia Clínica, Psicologia "
    "Hospitalar, Psicologia Escolar, Psicopedagogia, Psicologia Jurídica e demais áreas correlatas;")
add_clause_body(
    "b) MBA correlato à Psicologia, cobrindo Psicologia Organizacional e do Trabalho, Neurociência "
    "Executiva, Coaching, Liderança, Cultura Organizacional e Saúde Mental Corporativa em modalidade B2B;")
add_clause_body(
    "c) Mestrado Profissional em Psicologia (pós-graduação stricto sensu profissional regulamentada "
    "pela CAPES);")
add_clause_body(
    "d) Especialização Clínica certificada por Conselhos profissionais e entidades setoriais, "
    "incluindo, sem limitação, CFP (Conselho Federal de Psicologia, com observância à Resolução "
    "CFP 31/2022 e ao SATEPSI quando aplicável à Avaliação Psicológica), ABRAP, FBT, ABPp e EMDR Brasil;")
add_clause_body(
    "e) Formações híbridas e residências em Psicologia, com supervisão clínica regulamentada e "
    "eventuais parcerias institucionais com hospitais, laboratórios ou centros de pesquisa.")
add_clause_body(
    "1.3. O conjunto de atividades contratadas compreende, de forma integrada e aplicável a todas "
    "as cinco modalidades descritas na Cláusula 1.2:")
add_clause_body(
    "a) Execução técnica e operacional de GEO, incluindo auditoria de descobribilidade, otimização "
    "de Schema.org (com cobertura dos quatro tipos canônicos de programType — Specialization, MBA, "
    "ProfessionalMastersProgram e ProfessionalCertification —, além das estruturas Course e "
    "EducationalOccupationalProgram), llms.txt, robots.txt, sitemaps, internal linking e marcações "
    "de autoridade institucional (EducationalOrganization, hasCredential, Person para corpo docente);")
add_clause_body(
    "b) Geração de conteúdo editorial em padrão de qualidade compatível com publicações de "
    "referência (HBR/HSM/MIT Sloan), assinado ou referenciado conforme alinhamento entre as Partes, "
    "com cobertura editorial balanceada entre as cinco modalidades canônicas;")
add_clause_body(
    "c) Consultoria estratégica sobre presença em motores generativos, mention rate por persona e "
    "por modalidade, share-of-voice cross-LLM em cinco clusters semânticos (um por modalidade) "
    "frente a concorrentes diretos, e calibração da jornada de matrícula em pós-graduação em "
    "Psicologia em qualquer das modalidades;")
add_clause_body(
    "d) Assessoria executiva ao patrocinador executivo e à liderança de marketing da CONTRATANTE, "
    "com cadência de reuniões periódicas de status, decomposição de métricas por modalidade e "
    "recalibração mensal;")
add_clause_body(
    "e) Treinamento do time de marketing da CONTRATANTE em fundamentos e práticas de GEO, "
    "prompts-âncora canônicos (kit balanceado entre as cinco modalidades), engenharia de citação, "
    "padrão editorial HBR-grade e Voice Guard.")
add_clause_body(
    "1.4. O detalhamento operacional do escopo, incluindo entregáveis, métricas e cronograma, é o "
    "constante do roadmap do programa GEO IPOG mantido em repositório próprio sob governança da "
    "CONTRATADA, podendo ser ajustado de comum acordo entre as Partes ao longo da vigência. A "
    "inclusão, pela CONTRATANTE, de novos produtos de pós-graduação em Psicologia em qualquer das "
    "cinco modalidades canônicas ao longo da vigência integra automaticamente o escopo do programa, "
    "sem cobrança adicional, observado o limite de capacidade operacional descrito no roadmap.")

# ========== CLÁUSULA 2ª — PRAZO E VIGÊNCIA ==========
add_clause_title("2ª", "PRAZO E VIGÊNCIA")
add_clause_body(
    "2.1. O presente Contrato vigorará pelo prazo de 4 (quatro) meses consecutivos, contados da "
    "data de sua assinatura por ambas as Partes.")
add_clause_body(
    "2.2. Encerrado o prazo previsto na Cláusula 2.1, o Contrato extinguir-se-á automaticamente, "
    "sem necessidade de notificação prévia, ressalvada a possibilidade de prorrogação ou renovação "
    "mediante novo instrumento ou aditivo firmado pelas Partes.")

# ========== CLÁUSULA 3ª — VALOR E FORMA DE PAGAMENTO ==========
add_clause_title("3ª", "VALOR E FORMA DE PAGAMENTO")
add_clause_body(
    "3.1. Pela integralidade dos serviços descritos na Cláusula 1ª — abrangendo as cinco "
    "modalidades canônicas de forma integrada e indivisível — a CONTRATANTE pagará à CONTRATADA o "
    "valor mensal fixo de R$ 17.000,00 (dezessete mil reais), totalizando, ao final dos 4 (quatro) "
    "meses de vigência, o montante de R$ 68.000,00 (sessenta e oito mil reais).")
add_clause_body(
    "3.2. O valor mensal contempla, de forma consolidada e indivisível, todas as frentes descritas "
    "no objeto (execução, geração de conteúdo, consultoria, assessoria e treinamento) aplicadas a "
    "todas as cinco modalidades canônicas de pós-graduação em Psicologia da CONTRATANTE, sem "
    "cobranças adicionais por entregável, hora técnica, reunião, peça de conteúdo ou modalidade "
    "adicional incluída no portfólio da CONTRATANTE durante a vigência, salvo escopo expressamente "
    "fora do previsto neste Contrato e formalizado por aditivo.")
add_clause_body(
    "3.3. O pagamento será efetuado mensalmente, em até 10 (dez) dias corridos contados do "
    "recebimento, pela CONTRATANTE, da respectiva Nota Fiscal de Serviços emitida pela CONTRATADA, "
    "mediante depósito ou transferência (TED/PIX) na conta indicada abaixo:")
add_clause_body("Favorecida: BRGEO LTDA")
add_clause_body("CNPJ: 66.051.295/0001-33")
add_clause_body("Banco / Agência / Conta: Itaú Personalité, Agência 9049, Conta-Corrente 02757-7")
add_clause_body("Chave PIX: (62) 99814-1505")
add_clause_body(
    "3.4. Os tributos incidentes sobre os serviços, na forma da legislação aplicável, serão "
    "suportados pela parte legalmente responsável. Eventuais retenções na fonte por obrigação "
    "legal da CONTRATANTE serão informadas previamente e refletidas na Nota Fiscal.")
add_clause_body(
    "3.5. O atraso no pagamento de qualquer parcela acarretará multa moratória de 2% (dois por "
    "cento) sobre o valor em atraso, acrescida de juros de 1% (um por cento) ao mês, calculados "
    "pro rata die, sem prejuízo de correção monetária pelo IPCA/IBGE.")

# ========== CLÁUSULA 4ª — OBRIGAÇÕES DA CONTRATADA ==========
add_clause_title("4ª", "OBRIGAÇÕES DA CONTRATADA")
add_clause_body("4.1. A CONTRATADA se obriga a:")
add_clause_body(
    "a) Executar os serviços com diligência técnica e profissional, observando as boas práticas "
    "de mercado em GEO, SEO técnico, marketing de conteúdo e conformidade regulatória educacional "
    "aplicável a cada modalidade (MEC para Lato Sensu e MBA, CAPES para Mestrado Profissional, "
    "Conselhos profissionais para Especialização Clínica certificada);")
add_clause_body(
    "b) Disponibilizar Alexandre Caramaschi como interlocutor técnico-executivo principal e Head "
    "do programa, sem prejuízo do apoio de sua equipe e ferramentas;")
add_clause_body(
    "c) Manter cadência de reuniões mensais de status executivo com o CEO da CONTRATANTE e "
    "cadência operacional regular com a área de marketing, conforme alinhamento entre as Partes;")
add_clause_body(
    "d) Entregar relatórios periódicos consolidando métricas de mention rate, share-of-voice "
    "(decomposto pelos cinco clusters semânticos correspondentes às cinco modalidades), score "
    "técnico de Schema, progresso editorial e velocidade de fechamento de gaps, no formato e "
    "periodicidade combinados;")
add_clause_body(
    "e) Conduzir as ações de treinamento do time de marketing da CONTRATANTE em modalidade "
    "combinada entre as Partes (presencial em Goiânia, online ou híbrida);")
add_clause_body(
    "f) Tratar com confidencialidade toda informação não pública a que tiver acesso em razão deste "
    "Contrato, nos termos da Cláusula 7ª;")
add_clause_body(
    "g) Manter aderência ao naming canônico estabelecido para o programa, distinguindo "
    "expressamente o produto MBA da modalidade Especialização Lato Sensu, do Mestrado Profissional "
    "e da Especialização Clínica certificada em toda comunicação editorial e em toda marcação "
    "técnica de Schema.org.")

# ========== CLÁUSULA 5ª — OBRIGAÇÕES DA CONTRATANTE ==========
add_clause_title("5ª", "OBRIGAÇÕES DA CONTRATANTE")
add_clause_body("5.1. A CONTRATANTE se obriga a:")
add_clause_body(
    "a) Disponibilizar tempestivamente as informações, acessos, materiais institucionais e "
    "aprovações necessárias à execução dos serviços, incluindo o portfólio atualizado de "
    "pós-graduação em Psicologia em todas as modalidades vigentes;")
add_clause_body(
    "b) Designar um interlocutor operacional principal, sem prejuízo da relação executiva direta "
    "com o CEO Ronan Maia, com eventuais substituições comunicadas por escrito;")
add_clause_body("c) Efetuar os pagamentos nos prazos previstos na Cláusula 3ª;")
add_clause_body(
    "d) Aprovar previamente publicações ou mutações relevantes em propriedades digitais de sua "
    "titularidade, na forma do fluxo operacional combinado entre as Partes;")
add_clause_body(
    "e) Tomar, em tempo hábil, as decisões executivas sobre nomenclatura institucional, política "
    "de preços e portfólio por modalidade que sejam pré-requisito para a execução das frentes "
    "editorial e técnica do programa, conforme catalogadas no roadmap oficial.")

# ========== CLÁUSULA 6ª — PROPRIEDADE INTELECTUAL ==========
add_clause_title("6ª", "PROPRIEDADE INTELECTUAL")
add_clause_body(
    "6.1. Os entregáveis específicos produzidos pela CONTRATADA em nome da CONTRATANTE no âmbito "
    "deste Contrato — incluindo conteúdos editoriais finalizados, briefings entregues, dashboards "
    "customizados e auditorias técnicas — terão sua titularidade transferida à CONTRATANTE após o "
    "respectivo pagamento da parcela mensal correspondente.")
add_clause_body(
    "6.2. Permanecem de titularidade exclusiva da CONTRATADA os frameworks, metodologias, prompts "
    "canônicos, kits de prompt-âncora (incluindo o kit balanceado entre as cinco modalidades), "
    "padrões editoriais, voice guards, scripts, ferramentas, automações e quaisquer ativos de "
    "natureza genérica ou pré-existente utilizados na execução do Contrato, sendo concedida à "
    "CONTRATANTE licença de uso interno e não exclusivo enquanto vigente este instrumento.")
add_clause_body(
    "6.3. A CONTRATADA poderá mencionar publicamente, para fins de portfólio e referência "
    "comercial, o fato de ter prestado serviços à CONTRATANTE, vedada a divulgação de informações "
    "confidenciais ou de resultados quantitativos sem prévio alinhamento.")

# ========== CLÁUSULA 7ª — CONFIDENCIALIDADE E PROTEÇÃO DE DADOS ==========
add_clause_title("7ª", "CONFIDENCIALIDADE E PROTEÇÃO DE DADOS")
add_clause_body(
    "7.1. As Partes obrigam-se a manter sigilo sobre informações confidenciais, estratégicas, "
    "financeiras, comerciais, técnicas ou operacionais a que tiverem acesso em razão deste "
    "Contrato, abstendo-se de divulgá-las a terceiros sem prévia autorização escrita da outra "
    "Parte, durante toda a vigência e pelo prazo de 2 (dois) anos após o seu término.")
add_clause_body(
    "7.2. Não se consideram confidenciais informações que: (a) sejam ou se tornem de domínio "
    "público sem culpa da Parte receptora; (b) já fossem conhecidas pela Parte receptora antes da "
    "celebração deste Contrato; (c) devam ser divulgadas por força de lei ou decisão judicial.")
add_clause_body(
    "7.3. As Partes comprometem-se a observar a legislação aplicável de proteção de dados "
    "pessoais, em especial a Lei nº 13.709/2018 (LGPD), tratando os dados pessoais eventualmente "
    "compartilhados com a finalidade exclusiva da execução deste Contrato.")

# ========== CLÁUSULA 8ª — RESCISÃO ==========
add_clause_title("8ª", "RESCISÃO")
add_clause_body("8.1. Este Contrato poderá ser rescindido:")
add_clause_body(
    "a) Por comum acordo entre as Partes, a qualquer tempo, mediante termo de distrato;")
add_clause_body(
    "b) Por iniciativa unilateral de qualquer das Partes, mediante aviso prévio por escrito de 30 "
    "(trinta) dias, hipótese em que serão devidos à CONTRATADA os valores referentes aos serviços "
    "efetivamente prestados até a data de encerramento, incluído o período do aviso prévio;")
add_clause_body(
    "c) Por descumprimento contratual de qualquer das Partes, não sanado no prazo de 15 (quinze) "
    "dias corridos contados de notificação escrita da Parte adimplente.")
add_clause_body(
    "8.2. Em qualquer hipótese de rescisão, as Partes envidarão esforços razoáveis para conduzir "
    "transição ordenada das atividades em curso, sem prejuízo das obrigações de confidencialidade "
    "previstas na Cláusula 7ª.")

# ========== CLÁUSULA 9ª — NATUREZA DA RELAÇÃO ==========
add_clause_title("9ª", "NATUREZA DA RELAÇÃO")
add_clause_body(
    "9.1. O presente Contrato é firmado entre pessoas jurídicas em regime estritamente civil e "
    "comercial, não gerando vínculo empregatício, societário ou de exclusividade entre as Partes "
    "ou seus respectivos colaboradores.")
add_clause_body(
    "9.2. Cada Parte é integralmente responsável pelos encargos trabalhistas, previdenciários, "
    "fiscais e tributários relativos ao seu próprio pessoal.")

# ========== CLÁUSULA 10ª — DISPOSIÇÕES GERAIS ==========
add_clause_title("10ª", "DISPOSIÇÕES GERAIS")
add_clause_body(
    "10.1. As comunicações formais entre as Partes serão consideradas válidas quando realizadas "
    "por e-mail aos endereços indicados pelos respectivos representantes na execução do Contrato.")
add_clause_body(
    "10.2. A tolerância de qualquer das Partes quanto ao descumprimento de obrigações pela outra "
    "não importará em novação ou renúncia ao direito de exigir o seu fiel cumprimento.")
add_clause_body(
    "10.3. Caso qualquer cláusula deste Contrato venha a ser declarada nula ou inexequível, as "
    "demais permanecerão em pleno vigor.")
add_clause_body(
    "10.4. Eventuais alterações ao presente Contrato somente terão validade se formalizadas por "
    "termo aditivo assinado por ambas as Partes.")
add_clause_body(
    "10.5. A presente reemissão de 12 de maio de 2026 substitui integralmente, para todos os "
    "efeitos, a versão anterior do Contrato que tratava do recorte vertical de MBA Online de "
    "Psicologia. Não há ruptura de vigência: a contagem do prazo previsto na Cláusula 2.1 e os "
    "pagamentos previstos na Cláusula 3ª permanecem inalterados.")

# ========== CLÁUSULA 11ª — FORO ==========
add_clause_title("11ª", "FORO")
add_clause_body(
    "11.1. As Partes elegem o Foro da Comarca de Goiânia, Estado de Goiás, como competente para "
    "dirimir quaisquer dúvidas ou litígios decorrentes do presente Contrato, com renúncia "
    "expressa a qualquer outro, por mais privilegiado que seja.")

# ========== FECHAMENTO ==========
add_paragraph(
    "E, por estarem assim, justas e contratadas, as Partes assinam o presente Contrato em via "
    "eletrônica, com a mesma validade jurídica do instrumento físico, na forma da legislação aplicável.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12, space_before=18, space_after=12, line_spacing=1.5)
add_paragraph(
    "Goiânia/GO, _______________________ de _______________________ de 2026.",
    align=WD_ALIGN_PARAGRAPH.LEFT, size=12, space_before=6, space_after=24, line_spacing=1.5)

# ========== ASSINATURAS ==========
add_paragraph("CONTRATANTE", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
              size=12, space_before=6, space_after=4, line_spacing=1.2)
add_signature_line("IPOG — INSTITUTO DE PÓS-GRADUAÇÃO & GRADUAÇÃO LTDA")
add_paragraph("Representante: Ronan Maia", center=True, size=11, space_after=2, line_spacing=1.15)
add_paragraph("Cargo: Chief Executive Officer (CEO)", center=True, size=11, space_after=2, line_spacing=1.15)
add_paragraph("CPF: _______________________", center=True, size=11, space_after=2, line_spacing=1.15)
add_paragraph("CNPJ: 04.688.977/0001-02", center=True, size=11, space_after=14, line_spacing=1.15)

add_paragraph("CONTRATADA", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
              size=12, space_before=6, space_after=4, line_spacing=1.2)
add_signature_line("BRGEO LTDA")
add_paragraph("Representante: Alexandre Caramaschi", center=True, size=11, space_after=2, line_spacing=1.15)
add_paragraph("Cargo: Sócio-Administrador / CEO", center=True, size=11, space_after=2, line_spacing=1.15)
add_paragraph("CNPJ: 66.051.295/0001-33", center=True, size=11, space_after=18, line_spacing=1.15)

# ========== TESTEMUNHAS ==========
add_paragraph("TESTEMUNHAS", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT,
              size=12, space_before=12, space_after=12, line_spacing=1.2)

tbl = doc.add_table(rows=4, cols=3)
tbl.style = "Table Grid"
tbl.autofit = False
widths = [Cm(3.0), Cm(6.0), Cm(6.0)]
for row in tbl.rows:
    for idx, cell in enumerate(row.cells):
        cell.width = widths[idx]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

headers = ["", "Testemunha 1", "Testemunha 2"]
for idx, txt in enumerate(headers):
    cell = tbl.rows[0].cells[idx]
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(txt)
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)

rows_label = ["Nome", "CPF", "Assinatura"]
for ridx, label in enumerate(rows_label, start=1):
    cell_label = tbl.rows[ridx].cells[0]
    cell_label.text = ""
    p = cell_label.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(label)
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)
    for cidx in (1, 2):
        cell = tbl.rows[ridx].cells[cidx]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(" ")
        r.font.size = Pt(11)

# Salvar
OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT_PATH)
print(f"OK -> {OUT_PATH}")
print(f"Tamanho: {OUT_PATH.stat().st_size:,} bytes")
